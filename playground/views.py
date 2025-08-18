from django.shortcuts import render, redirect
from .forms import PlaygroundModelForm
from .models import PlaygroundModel
import csv
from django.http import HttpResponse
from django.core.exceptions import ValidationError

import csv
from io import TextIOWrapper


def home(request):
    return render(request, 'playground_1.html')

def playground_table(request):
    phn_query = request.GET.get('phn', '')
    playground_data = []
    if phn_query:
        playground_data = PlaygroundModel.objects.filter(PHN=phn_query).order_by('PHN')
    return render(request, 'playground_table.html', {
        'playground_data': playground_data,
        'phn_query': phn_query
    })
    
def download_phn(request):
    from docx import Document
    from django.utils.text import slugify
    from io import BytesIO
    if request.method == 'GET':
        return render(request, 'download_phn_form.html')

    # POST: generate document
    name = request.POST.get('name')
    phil_no = request.POST.get('phil_no')
    date_hired = request.POST.get('date_hired')
    position = request.POST.get('position')
    phn_query = request.GET.get('phn', '')
    query = request.GET.get('q', '')
    # Filter data as in all_data_table
    if query:
        data = PlaygroundModel.objects.filter(
            surname__icontains=query
        ) | PlaygroundModel.objects.filter(
            given_name__icontains=query
        ) | PlaygroundModel.objects.filter(
            PHN__icontains=query
        ) | PlaygroundModel.objects.filter(
            year__icontains=query
        ) | PlaygroundModel.objects.filter(
            month__icontains=query
        )
    elif phn_query:
        data = PlaygroundModel.objects.filter(PHN=phn_query).order_by('PHN')
    else:
        data = PlaygroundModel.objects.all().order_by('PHN')

    doc = Document()
    from datetime import datetime
    today_str = datetime.now().strftime('%B %d, %Y')
    # Letter-style spacing
    doc.add_paragraph(today_str)
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_heading("CERTIFICATION", level=1)
    doc.add_paragraph("")
    doc.add_paragraph(
        f"This is to certify that {name} with PhilHealth No. {phil_no} is an employee of Carelon Global Solutions Philippines, Inc. (formerly known as Legato Health Technologies Philippines, Inc) since {date_hired} and is currently holding the position title of {position}."
    )
    doc.add_paragraph("")
    doc.add_paragraph("The following are the latest PhilHealth contributions.")
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'OR#'
    hdr_cells[1].text = 'EE SHARE'
    hdr_cells[2].text = 'ER SHARE'
    hdr_cells[3].text = 'TOTAL'
    hdr_cells[4].text = 'DATE REMITTED'
    hdr_cells[5].text = 'APPLICABLE MONTH'
    # Example: fill with data from PlaygroundModel or static rows
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.PHN)  # Replace with OR# if available
        ps_val = float(item.ps.to_decimal()) if item.ps is not None else 0
        es_val = float(item.es.to_decimal()) if item.es is not None else 0
        row_cells[1].text = str(ps_val) if item.ps is not None else ''
        row_cells[2].text = str(es_val) if item.es is not None else ''
        row_cells[3].text = str(ps_val + es_val)
        row_cells[4].text = str(item.year)  # Replace with remittance date if available
        row_cells[5].text = str(item.month)
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Noted by:")
    doc.add_paragraph("")
    doc.add_paragraph("ERIENNE F. MACEDA")
    doc.add_paragraph("HR Consultant")
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    filename = f"certification_{slugify(name or 'employee')}.docx"
    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def add_playground(request):
    message = None
    message_type = None
    if request.method == 'POST':
        if request.FILES.get('csv_file'):
            csv_file = TextIOWrapper(request.FILES['csv_file'].file, encoding='utf-8')
            reader = csv.DictReader(csv_file)
            success_count = 0
            error_count = 0
            for row in reader:
                def clean_decimal(val):
                    if val:
                        return val.replace(',', '')
                    return None
                try:
                    PlaygroundModel.objects.create(
                        month=int(row.get('month')) if row.get('month') else None,
                        year=int(row.get('year')) if row.get('year') else None,
                        PHN=row.get('PHN'),
                        surname=row.get('surname') or '',
                        given_name=row.get('given_name') or '',
                        middle_name=row.get('middle_name') or '',
                        ps=clean_decimal(row.get('ps')),
                        es=clean_decimal(row.get('es')),
                        status=row.get('status') or '',
                    )
                    success_count += 1
                except (ValueError, ValidationError):
                    error_count += 1
                    continue
            if error_count == 0 and success_count > 0:
                message = f"Upload complete! {success_count} records added."
                message_type = "success"
            elif success_count > 0:
                message = f"Upload finished: {success_count} records added, {error_count} rows skipped due to errors."
                message_type = "warning"
            else:
                message = "Upload failed: All rows had errors."
                message_type = "error"
        else:
            message = "No file uploaded."
            message_type = "error"
        return render(request, 'add_playground.html', {'message': message, 'message_type': message_type})
    return render(request, 'add_playground.html')

def all_data_table(request):
    query = request.GET.get('q', '')
    if query:
        playground_data = PlaygroundModel.objects.filter(
            surname__icontains=query
        ) | PlaygroundModel.objects.filter(
            given_name__icontains=query
        ) | PlaygroundModel.objects.filter(
            PHN__icontains=query
        ) | PlaygroundModel.objects.filter(
            year__icontains=query
        ) | PlaygroundModel.objects.filter(
            month__icontains=query
        )
    else:
        playground_data = PlaygroundModel.objects.all().order_by('PHN')
    return render(request, 'all_data_table.html', {
        'playground_data': playground_data,
        'query': query
    })